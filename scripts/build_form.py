# -*- coding: utf-8 -*-
"""
质量管理现场表单生成器（下一级执行文件：检验记录表 / 试验测试报告表 / 点检表）

产出可直接打印下发的空白模板：
  - Excel (.xlsx)：带表头信息区、字段表头、空白数据行、判定/结果下拉、签字栏、打印设置
  - Word (.docx)：同结构

用法：
  python build_form.py                                      # 内置样本，产出三类演示模板（xlsx+docx）
  python build_form.py --type inspection --format both --rows 15 \
      --title "XX首件检验记录表" --meta "产品名称=XX" "零件号=ABC-001"
  python build_form.py --type test --format xlsx
  python build_form.py --type checklist --format docx

依赖：openpyxl, python-docx   （pip install openpyxl python-docx）
"""
import argparse
import os
from datetime import date

# 主色（企业色）
PRIMARY_RGB = "C8102E"

# ---- 三类表单定义 ----
FORMS = {
    "inspection": {
        "label": "检验记录表",
        "default_title": "产品检验记录表",
        "meta": ["表单名称", "产品名称/型号", "零件号/图号", "批次/订单号",
                 "检验阶段(首检/巡检/终检)", "检验依据(SIP/标准)", "日期", "班次"],
        "columns": ["序号", "检验项目", "标准要求", "实测值", "单位",
                    "判定(合格/不合格)", "不良现象/备注", "检验员"],
        "dropdown_col": "判定(合格/不合格)",
        "dropdown": ["合格", "不合格"],
        "signoff": ["检验员", "审核", "日期"],
    },
    "test": {
        "label": "试验/测试报告表",
        "default_title": "试验测试报告表",
        "meta": ["表单名称", "试验名称", "样品编号", "样品状态",
                 "试验设备", "环境(温湿度)", "依据标准", "日期"],
        "columns": ["序号", "试验项目", "试验条件/方法", "技术要求", "测试结果",
                    "单位", "结论(合格/不合格)", "偏离说明", "测试人", "复核人"],
        "dropdown_col": "结论(合格/不合格)",
        "dropdown": ["合格", "不合格"],
        "signoff": ["测试人", "复核人", "日期"],
    },
    "checklist": {
        "label": "点检表",
        "default_title": "设备点检表",
        "meta": ["表单名称", "设备名称", "设备编号", "点检部位/区域",
                 "点检周期(班/日/周/月)", "日期", "班次"],
        "columns": ["序号", "点检项目", "点检标准/要求", "点检方法",
                    "结果(正常/异常)", "异常处理/措施", "点检人", "确认人"],
        "dropdown_col": "结果(正常/异常)",
        "dropdown": ["正常", "异常"],
        "signoff": ["点检人", "确认人", "日期"],
    },
}

SAMPLE_META = {
    "inspection": {"表单名称": "首件检验记录表", "产品名称/型号": "六角头螺栓 M10×40",
                   "零件号/图号": "GB/T 5783", "批次/订单号": "B-2026-0714",
                   "检验阶段(首检/巡检/终检)": "首检", "检验依据(SIP/标准)": "SIP-QI-001",
                   "日期": str(date.today()), "班次": "早班"},
    "test": {"表单名称": "盐雾试验报告表", "试验名称": "中性盐雾试验", "样品编号": "S-001",
             "样品状态": "合格件", "试验设备": "盐雾试验箱 YW-90", "环境(温湿度)": "23℃/50%RH",
             "依据标准": "GB/T 10125", "日期": str(date.today())},
    "checklist": {"表单名称": "日常点检表", "设备名称": "数控车床", "设备编号": "CNC-03",
                  "点检部位/区域": "主轴/液压/润滑", "点检周期(班/日/周/月)": "班",
                  "日期": str(date.today()), "班次": "早班"},
}

# 列宽预设（按最大列数兜底）
WIDTHS = [6, 18, 22, 14, 10, 16, 20, 12, 12, 12]


def _col_letter(i):
    from openpyxl.utils import get_column_letter
    return get_column_letter(i)


def build_xlsx(form, meta, n_rows, title, out_path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation
    from openpyxl.worksheet.properties import PageSetupProperties
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    # sheet 名不能含 : \ / ? * [ ]，且最长 31 字符
    ws.title = form["label"].replace("/", "-")[:31]
    thin = Side(style="thin", color="BBBBBB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    fill_pri = PatternFill("solid", fgColor=PRIMARY_RGB)
    fill_gray = PatternFill("solid", fgColor="F2F2F2")
    font_white = Font(color="FFFFFF", bold=True, size=11)
    font_title = Font(color="FFFFFF", bold=True, size=14)
    font_bold = Font(bold=True)
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ncol = len(form["columns"])
    last_col = get_column_letter(ncol)

    # 标题
    ws.merge_cells(f"A1:{last_col}1")
    c = ws["A1"]
    c.value = title
    c.fill = fill_pri
    c.font = font_title
    c.alignment = align_center
    ws.row_dimensions[1].height = 28

    # 表头信息区
    r = 2
    for k in form["meta"]:
        a = ws.cell(r, 1, k)
        a.font = font_bold
        a.fill = fill_gray
        a.border = border
        a.alignment = align_left
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=ncol)
        b = ws.cell(r, 2, meta.get(k, "【待填写】"))
        b.border = border
        b.alignment = align_left
        r += 1

    r += 1  # 空行

    # 字段表头
    hdr_row = r
    for i, col in enumerate(form["columns"], start=1):
        cell = ws.cell(hdr_row, i, col)
        cell.fill = fill_pri
        cell.font = font_white
        cell.border = border
        cell.alignment = align_center
    r += 1

    # 空白数据行
    for _ in range(n_rows):
        for i in range(1, ncol + 1):
            cell = ws.cell(r, i)
            cell.border = border
            cell.alignment = align_center
        r += 1

    # 判定/结果下拉
    if form.get("dropdown_col") and form["dropdown_col"] in form["columns"]:
        col_idx = form["columns"].index(form["dropdown_col"]) + 1
        L = get_column_letter(col_idx)
        dv = DataValidation(type="list",
                            formula1='"' + ",".join(form["dropdown"]) + '"',
                            allow_blank=True)
        ws.add_data_validation(dv)
        dv.add(f"{L}{hdr_row + 1}:{L}{hdr_row + n_rows}")

    # 签字栏
    r += 1
    for k in form["signoff"]:
        a = ws.cell(r, 1, k)
        a.font = font_bold
        a.border = border
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=ncol)
        ws.cell(r, 2).border = border
        r += 1

    # 列宽 & 冻结
    for i in range(1, ncol + 1):
        w = WIDTHS[i - 1] if i - 1 < len(WIDTHS) else 14
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{hdr_row + 1}"

    # 打印：横向、适应一页宽
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    try:
        ws.sheet_properties.pageSetUpPr = PageSetupProperties(fitToPage=True)
    except Exception:
        pass

    wb.save(out_path)


def build_docx(form, meta, n_rows, title, out_path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hexcolor)
        tcPr.append(shd)

    doc = Document()
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)

    # 表头信息区（2列）
    meta_tbl = doc.add_table(rows=0, cols=2)
    meta_tbl.style = "Table Grid"
    for k in form["meta"]:
        row = meta_tbl.add_row().cells
        row[0].text = k
        row[1].text = meta.get(k, "【待填写】")
        shade(row[0], "F2F2F2")
    doc.add_paragraph("")

    # 主表
    t = doc.add_table(rows=1, cols=len(form["columns"]))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, col in enumerate(form["columns"]):
        hdr[i].text = col
        shade(hdr[i], PRIMARY_RGB)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for _ in range(n_rows):
        t.add_row().cells

    doc.add_paragraph("")
    sign = doc.add_paragraph()
    sign.add_run("　".join(f"{k}：____________" for k in form["signoff"]))

    note = doc.add_paragraph()
    nr = note.add_run("★待确认：检验项目、标准、抽样方案由使用者按图纸/客户/SIP 提供并确认后使用。")
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="质量管理现场表单生成器（Excel/Word）")
    parser.add_argument("--type", choices=list(FORMS.keys()), default="",
                        help="表单类型：inspection / test / checklist，缺省产出全部三类")
    parser.add_argument("--format", choices=["xlsx", "docx", "both"], default="both")
    parser.add_argument("--title", default="", help="表单标题，覆盖默认")
    parser.add_argument("--rows", type=int, default=15, help="空白数据行数，默认 15")
    parser.add_argument("--meta", nargs="*", default=[], help="表头字段 key=value（可重复）")
    parser.add_argument("--out-dir", default="", help="输出目录，默认当前工作目录（用户运行技能处）")
    args = parser.parse_args()

    types = [args.type] if args.type else list(FORMS.keys())
    script_dir = os.path.dirname(os.path.abspath(__file__))
    # 产物默认写到用户当前工作目录（技能运行处），不污染技能目录本身
    out_dir = args.out_dir or os.getcwd()
    os.makedirs(out_dir, exist_ok=True)

    for t in types:
        form = FORMS[t]
        meta = dict(SAMPLE_META.get(t, {}))
        for kv in args.meta:
            if "=" in kv:
                k, v = kv.split("=", 1)
                meta[k] = v
        title = args.title or form["default_title"]
        base = f"form_{t}"
        if args.format in ("xlsx", "both"):
            p = os.path.join(out_dir, base + ".xlsx")
            build_xlsx(form, meta, args.rows, title, p)
            print(f"[OK] XLSX -> {os.path.abspath(p)}")
        if args.format in ("docx", "both"):
            p = os.path.join(out_dir, base + ".docx")
            build_docx(form, meta, args.rows, title, p)
            print(f"[OK] DOCX -> {os.path.abspath(p)}")


if __name__ == "__main__":
    main()
